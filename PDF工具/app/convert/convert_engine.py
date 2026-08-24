import fitz
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io


class CancelledException(Exception):
    """操作被用户取消"""
    pass


class ConvertEngine:
    """PDF 转 DOCX 引擎"""

    def __init__(self, progress_callback=None, cancel_check=None):
        """
        progress_callback: function(current, total, message) -> None
        cancel_check: function() -> bool, 返回 True 表示已取消
        """
        self._progress_callback = progress_callback
        self._cancel_check = cancel_check

    def _check_cancelled(self):
        if self._cancel_check and self._cancel_check():
            raise CancelledException("操作已取消")

    def _report(self, current: int, total: int, message: str = ""):
        self._check_cancelled()
        if self._progress_callback:
            self._progress_callback(current, total, message)

    def convert(self, input_path: str, output_path: str, dpi: int = 150) -> str:
        if not os.path.isfile(input_path):
            raise FileNotFoundError(f"文件不存在: {input_path}")

        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        doc = fitz.open(input_path)
        total_pages = doc.page_count

        output_docx = Document()
        style = output_docx.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

        self._report(0, total_pages, "开始分析文档...")

        has_text = False
        for page_num in range(total_pages):
            self._check_cancelled()
            page = doc[page_num]
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block["type"] == 0:
                    for line in block["lines"]:
                        for span in line["spans"]:
                            if span["text"].strip():
                                has_text = True
                                break
                        if has_text:
                            break
                    if has_text:
                        break
            if has_text:
                break

        if has_text:
            result = self._convert_text_based(doc, output_docx, output_path, total_pages)
        else:
            result = self._convert_image_based(doc, output_docx, output_path, total_pages, dpi)

        doc.close()
        return result

    def _add_rich_paragraph(self, output_docx, lines_data: list) -> None:
        """
        将一组 line 数据写入一个 DOCX 段落，保留字体/格式
        lines_data: 合并后的行数据，每行包含 spans 列表
        """
        if not lines_data:
            return

        p = output_docx.add_paragraph()

        # 对齐方式检测（基于首个 block 的坐标）
        first_line = lines_data[0]
        if "bbox" in first_line:
            x0 = first_line["bbox"][0]
            page_w = first_line.get("page_width", 612)
            # 居中: x0 在页面中部附近
            if abs(x0 - page_w / 2) < page_w * 0.15:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif x0 > page_w * 0.3:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 段间距
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15

        for line_data in lines_data:
            spans = line_data.get("spans", [])
            for span in spans:
                text = span.get("text", "")
                if not text.strip():
                    continue

                run = p.add_run(text)

                # 字体名称
                font_name = span.get("font", "")
                if font_name:
                    # 映射常见字体名
                    run.font.name = font_name
                    # 设置中文字体回退
                    r = run._element
                    rPr = r.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                    if rPr is None:
                        from lxml import etree
                        rPr = etree.SubElement(r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                    rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
                    if rFonts is None:
                        from lxml import etree
                        rFonts = etree.SubElement(rPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
                    rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', font_name)

                # 字号（pt）
                size = span.get("size", 11)
                run.font.size = Pt(size)

                # 粗体（flags 第 5 位）
                flags = span.get("flags", 0)
                run.font.bold = bool(flags & 32)  # 0x20

                # 斜体（flags 第 1 位）
                run.font.italic = bool(flags & 2)  # 0x02

                # 颜色
                color = span.get("color", 0)
                if color and color != 0:
                    r = (color >> 16) & 0xFF
                    g = (color >> 8) & 0xFF
                    b = color & 0xFF
                    run.font.color.rgb = RGBColor(r, g, b)

    def _convert_text_based(self, doc, output_docx, output_path, total_pages):
        """文本型 PDF 转换（保留字体/字号/粗斜/对齐）"""
        self._report(0, total_pages, "正在提取文本与格式...")

        for page_num in range(total_pages):
            self._report(page_num + 1, total_pages, f"正在处理第 {page_num + 1} 页...")

            page = doc[page_num]
            page_rect = page.rect
            page_w = page_rect.width

            page_dict = page.get_text("dict")
            blocks = page_dict["blocks"]

            # 分离文本块和图片块
            text_blocks = [b for b in blocks if b["type"] == 0]
            image_blocks = [b for b in blocks if b["type"] == 1]

            # 按 y 坐标排序
            text_blocks.sort(key=lambda b: b["bbox"][1])

            # 合并相邻行到同一段落（y 坐标接近）
            paragraphs: list[list] = []
            current_lines = []
            last_y = -1

            for block in text_blocks:
                block_y = block["bbox"][1]
                lines = block.get("lines", [])

                for line in lines:
                    line_y = line["bbox"][1]
                    line["page_width"] = page_w

                    if last_y >= 0 and abs(line_y - last_y) > page_rect.height * 0.025:
                        if current_lines:
                            paragraphs.append(current_lines)
                            current_lines = []
                    current_lines.append(line)
                    last_y = line_y

            if current_lines:
                paragraphs.append(current_lines)

            # 写入 DOCX 段落
            for para_lines in paragraphs:
                self._check_cancelled()
                self._add_rich_paragraph(output_docx, para_lines)

            # 嵌入图片
            for block in image_blocks:
                try:
                    self._check_cancelled()
                    x0, y0, x1, y1 = block["bbox"]
                    clip = fitz.Rect(x0, y0, x1, y1)
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), clip=clip)
                    img_bytes = pix.tobytes("png")
                    img_stream = io.BytesIO(img_bytes)

                    p = output_docx.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_width = (x1 - x0) * 2
                    max_width = Cm(15)
                    if img_width > max_width:
                        p.add_run().add_picture(img_stream, width=max_width)
                    else:
                        p.add_run().add_picture(img_stream, width=Cm(img_width / 10))
                except Exception:
                    pass

            if page_num < total_pages - 1:
                output_docx.add_page_break()

        self._report(total_pages, total_pages, "正在保存 DOCX...")
        output_docx.save(output_path)

        abs_path = os.path.abspath(output_path)
        self._report(total_pages, total_pages, f"转换完成: {os.path.basename(abs_path)}")
        return abs_path

    def _convert_image_based(self, doc, output_docx, output_path, total_pages, dpi):
        """扫描型 PDF 转换（降级为图片嵌入）"""
        zoom = dpi / 72
        matrix = fitz.Matrix(zoom, zoom)

        self._report(0, total_pages, "扫描型 PDF，将以图片形式嵌入...")

        for page_num in range(total_pages):
            self._check_cancelled()
            self._report(page_num + 1, total_pages, f"正在渲染第 {page_num + 1} 页...")

            page = doc[page_num]
            pix = page.get_pixmap(matrix=matrix)
            img_bytes = pix.tobytes("png")
            img_stream = io.BytesIO(img_bytes)

            p = output_docx.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"第 {page_num + 1} 页（图片型）")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            output_docx.add_picture(img_stream, width=Inches(6))

            if page_num < total_pages - 1:
                output_docx.add_page_break()

        self._report(total_pages, total_pages, "正在保存 DOCX...")
        output_docx.save(output_path)

        abs_path = os.path.abspath(output_path)
        self._report(total_pages, total_pages, f"转换完成: {os.path.basename(abs_path)}")
        return abs_path