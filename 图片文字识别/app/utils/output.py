"""输出格式转换模块。

支持格式：
  - txt        纯文本
  - markdown   Markdown（按页/按行）
  - json       含坐标的 JSON
  - searchable_pdf  可搜索 PDF（原图 + 隐形文字层）
  - original   还原为原格式（图片→txt, PDF→可搜索PDF, DOCX→docx/txt）

输出文件名以 '_' 开头，确保排在目录顶部（用户偏好）。
"""
from __future__ import annotations

import json as _json
import logging
import os
from typing import List, Optional

from ..core.document.base import DocumentResult, PageResult
from .text_dehardwrap import dehard_wrap, merge_lines_to_paragraphs

logger = logging.getLogger(__name__)

# 全文级文本纠错器懒加载单例
_FULLTEXT_CORRECTOR = None
_FULLTEXT_CORRECTOR_LOCK = None


def _get_fulltext_corrector():
    """获取全文级纠错器（懒加载，与行级纠错器共用同一实例）。"""
    global _FULLTEXT_CORRECTOR, _FULLTEXT_CORRECTOR_LOCK
    if _FULLTEXT_CORRECTOR is None:
        try:
            import threading as _t
            if _FULLTEXT_CORRECTOR_LOCK is None:
                _FULLTEXT_CORRECTOR_LOCK = _t.Lock()
            with _FULLTEXT_CORRECTOR_LOCK:
                if _FULLTEXT_CORRECTOR is None:
                    from ..core.text_corrector import TextCorrector
                    _FULLTEXT_CORRECTOR = TextCorrector()
                    logger.info("全文级纠错器已初始化（输出模块懒加载）")
        except Exception as e:
            logger.warning("全文级纠错器初始化失败，跳过全文纠错: %s", e)
            _FULLTEXT_CORRECTOR = False
    return _FULLTEXT_CORRECTOR if _FULLTEXT_CORRECTOR is not False else None


def _apply_fulltext_correction(text: str, source_name: str = "") -> str:
    """对全文做最后一层文本纠错（段落级 + 串栏断裂修复）。

    参数:
        text: 原始全文本（按页用换行拼接后的大字符串）
        source_name: 用于日志标注的来源名（如文件名）
    返回:
        纠错后的全文
    """
    if not text:
        return text
    corrector = _get_fulltext_corrector()
    if corrector is None:
        return text
    try:
        # 先按换行拆成行（纠错器的 correct_lines 接收行列表），
        # 以与行级纠错相同的管道再次执行全文语境
        lines = text.split("\n")
        res = corrector.correct_lines(lines, merge_to_paragraphs=True)
        if res.total_fixes > 0:
            from ..core.text_corrector import TextCorrector
            log_str = TextCorrector.format_records(res.records, max_items=10)
            logger.info(
                "全文级纠错[%s]: 共修复 %d 处。\n%s",
                source_name or "未命名", res.total_fixes, log_str,
            )
            # 写入进度日志（便于用户在控制台直接看到）
            _log_writer(f"[全文纠错] {source_name}: 修复 {res.total_fixes} 处")
        return res.text
    except Exception as e:
        logger.warning("全文级纠错异常（跳过，不阻塞输出）: %s", e)
        return text


# 进度日志：直接写文件，绕开 logging 系统（uvicorn 会覆盖 logging 配置）
_PROGRESS_LOG: Optional[str] = None


def _log_writer(tag: str) -> None:
    """把 writer 相关日志追加到进度日志文件。"""
    global _PROGRESS_LOG
    if _PROGRESS_LOG is None:
        from .task_dirs import progress_log_path, LOG_DIR
        os.makedirs(LOG_DIR, exist_ok=True)
        _PROGRESS_LOG = progress_log_path()
    import time as _t
    ts = _t.strftime("%Y-%m-%d %H:%M:%S", _t.localtime())
    try:
        with open(_PROGRESS_LOG, "a", encoding="utf-8") as f:
            f.write(f"{ts} | {tag}\n")
    except Exception:
        pass


# ----------------------------------------------------------------------
# 输出路径
# ----------------------------------------------------------------------
def _output_path(
    source_path: str, ext: str, output_dir: str, suffix: str = "_OCR"
) -> str:
    """生成输出文件路径，文件名以 _ 开头。"""
    base = os.path.basename(source_path)
    name, _ = os.path.splitext(base)
    out_dir = output_dir if output_dir else os.path.dirname(source_path)
    out_name = f"_{name}{suffix}.{ext.lstrip('.')}"
    return os.path.join(out_dir, out_name)


# ----------------------------------------------------------------------
# 文本类输出
# ----------------------------------------------------------------------
def to_txt(result: DocumentResult) -> str:
    parts = []
    if result.native_text:
        # 原生文字（DOCX 段落）已是真实段落结构，保留原换行，不去硬换行
        parts.append(result.native_text)
    source_name = os.path.basename(result.source_path)
    for page in result.pages:
        if page.skipped:
            parts.append(f"[第 {page.page_no} 页/图: 跳过 - {page.reason}]")
            continue
        # 先做全文纠错（跨句 + 段落级修复），再去硬换行
        page_text = _apply_fulltext_correction(
            page.ocr_result.text,
            source_name=f"{source_name} - TXT p{page.page_no}",
        )
        parts.append(dehard_wrap(page_text))
    return "\n\n".join(p for p in parts if p)


def to_markdown(result: DocumentResult) -> str:
    lines = [f"# {os.path.basename(result.source_path)}", ""]
    if result.native_text:
        # 原生文字保留原段落结构，不去硬换行
        lines.append("## 原生文字\n")
        lines.append(result.native_text)
        lines.append("")
    source_name = os.path.basename(result.source_path)
    for page in result.pages:
        title = f"第 {page.page_no} 页" if len(result.pages) > 1 else "识别结果"
        if page.skipped:
            lines.append(f"## {title}\n\n> 跳过: {page.reason}\n")
            continue
        lines.append(f"## {title}\n")
        # 先纠错再去硬换行（仅对 OCR 文本，Markdown 标题行不会被合并：
        # 标题以 # 开头独占一行，且前后有空行分隔）
        page_text = _apply_fulltext_correction(
            page.ocr_result.text,
            source_name=f"{source_name} - MD p{page.page_no}",
        )
        lines.append(dehard_wrap(page_text))
        lines.append("")
    return "\n".join(lines)


def to_json(result: DocumentResult) -> str:
    data = {
        "source": result.source_path,
        "native_text": result.native_text,
        "pages": [
            {
                "page_no": p.page_no,
                "skipped": p.skipped,
                "reason": p.reason,
                "provider": p.ocr_result.provider,
                "lines": [
                    {
                        "text": ln.text,
                        "coords": ln.coords,
                        "confidence": ln.confidence,
                    }
                    for ln in p.ocr_result.lines
                ],
            }
            for p in result.pages
        ],
    }
    return _json.dumps(data, ensure_ascii=False, indent=2)


# ----------------------------------------------------------------------
# 可搜索 PDF（原图 + 隐形文字层）
# ----------------------------------------------------------------------
def to_searchable_pdf(result: DocumentResult, render_dpi: int = 200) -> bytes:
    """生成可搜索 PDF：每页插入原图，叠加隐形文字层。"""
    import fitz

    out_doc = fitz.open()
    scale = 72.0 / render_dpi  # 像素 -> PDF 点
    for page in result.pages:
        if page.image is None:
            continue
        h, w = page.image.shape[:2]
        # PDF 页面尺寸（点）
        rect = fitz.Rect(0, 0, w * scale, h * scale)
        pdf_page = out_doc.new_page(width=rect.width, height=rect.height)
        # 插入原图
        img_bytes = _cv2_to_png_bytes(page.image)
        pdf_page.insert_image(rect, stream=img_bytes)
        # 叠加隐形文字层
        if not page.skipped:
            _insert_invisible_text(pdf_page, page, scale)
    # PyMuPDF 1.28 不接受 bytearray，需用 BytesIO 文件对象
    import io
    buf = io.BytesIO()
    out_doc.save(buf)
    out_doc.close()
    return buf.getvalue()


# ----------------------------------------------------------------------
# 中文字体定位（避免打包后内置 CJK 字体不可用）
# ----------------------------------------------------------------------
_CJK_FONT_CACHE: Optional[tuple] = None  # (fontfile, fontname) 或 None


def _get_cjk_font() -> Optional[tuple]:
    """返回可用的中文字体 (fontfile, fontname)，找不到返回 None。

    查找优先级：
      1. 打包内置字体 _internal/fonts/（确保任意 Windows 可用，含英文版）
      2. Windows 系统字体 C:/Windows/Fonts/
      3. PyMuPDF 内置 china-s（开发环境可用，打包后可能丢失）
    """
    global _CJK_FONT_CACHE
    if _CJK_FONT_CACHE is not None:
        return _CJK_FONT_CACHE if _CJK_FONT_CACHE != ("", "") else None

    import os
    import sys

    # 打包内置字体目录（spec 中打包到 _internal/fonts/）
    bundled_font_dir = None
    if getattr(sys, "frozen", False):
        exe_dir = os.path.dirname(sys.executable)
        bundled_font_dir = os.path.join(exe_dir, "_internal", "fonts")
    else:
        # 开发环境：项目根目录下的 fonts/
        root = os.path.dirname(os.path.dirname(os.path.dirname(
            os.path.abspath(__file__))))
        bundled_font_dir = os.path.join(root, "fonts")

    # 字体候选（优先 .ttf 支持子集化，其次 .ttc）
    # 每项: (bundled_path, system_path, fontname)
    candidates = [
        ("simhei.ttf",  "C:/Windows/Fonts/simhei.ttf",  "simhei"),
        ("simkai.ttf",  "C:/Windows/Fonts/simkai.ttf",  "simkai"),
        ("simsun.ttc",  "C:/Windows/Fonts/simsun.ttc",  "simsun"),
        ("msyh.ttc",    "C:/Windows/Fonts/msyh.ttc",    "msyh"),
    ]
    for bundled_name, sys_path, fontname in candidates:
        # 1. 打包内置字体
        if bundled_font_dir:
            bundled_path = os.path.join(bundled_font_dir, bundled_name)
            if os.path.exists(bundled_path):
                _CJK_FONT_CACHE = (bundled_path, fontname)
                logger.info("中文字体(打包内置): %s (%s)", fontname, bundled_path)
                return _CJK_FONT_CACHE
        # 2. 系统字体
        if os.path.exists(sys_path):
            _CJK_FONT_CACHE = (sys_path, fontname)
            logger.info("中文字体(系统): %s (%s)", fontname, sys_path)
            return _CJK_FONT_CACHE

    # 备用：内置 china-s（打包后可能不可用，但开发环境可用）
    _CJK_FONT_CACHE = ("", "")
    logger.warning("未找到系统中文字体，回退到内置 china-s（打包后可能失效）")
    return None


def _insert_invisible_text(pdf_page, page: PageResult, scale: float) -> None:
    """在 PDF 页面插入隐形文字（可选中、可搜索，但不可见）。

    关键点：
    1. 用 Windows 系统字体文件（simsun.ttc），不依赖 PyMuPDF 内置 CJK 字体，
       确保 PyInstaller 打包后行为一致。
    2. 用 insert_text 而非 insert_textbox：textbox 当 rect 过小会丢弃文字，
       insert_text 直接按坐标写入，不受 rect 容量限制。
    3. render_mode=3 使文字不可见但可搜索可选中。
    """
    import fitz

    font = _get_cjk_font()
    font_kwargs = {}
    if font:
        font_kwargs["fontfile"] = font[0]
        font_kwargs["fontname"] = font[1]
    else:
        # 备用：内置 CJK 字体
        font_kwargs["fontname"] = "china-s"

    inserted = 0
    failed = 0
    for line in page.ocr_result.lines:
        if not line.coords or not line.text:
            continue
        xs = [p[0] for p in line.coords]
        ys = [p[1] for p in line.coords]
        # 坐标转换到 PDF 坐标系（原图像素 → PDF 点）
        x0, y0 = min(xs) * scale, min(ys) * scale
        x1, y1 = max(xs) * scale, max(ys) * scale
        # 文字高度（PDF 坐标系），至少 6 点保证可读性
        font_size = max((y1 - y0) * 0.9, 6)
        # insert_text 的 origin 是文字基线左下角
        # y1 是文字框底部（近似基线位置）
        try:
            pdf_page.insert_text(
                (x0, y1),  # 左下角作为基线起点
                line.text,
                fontsize=font_size,
                render_mode=3,  # 3 = 不可见（可选中、可搜索，但视觉不可见）
                **font_kwargs,
            )
            inserted += 1
        except Exception as e:
            failed += 1
            if failed <= 3:
                logger.warning("文字插入失败: %s, text=%r", e, line.text[:20])
    logger.info("文字层写入: 成功 %d / 失败 %d / 共 %d 行", inserted, failed, len(page.ocr_result.lines))


def fitz_rect(x0, y0, x1, y1):
    import fitz

    return fitz.Rect(x0, y0, x1, y1)


def _cv2_to_png_bytes(image) -> bytes:
    import cv2

    ok, buf = cv2.imencode(".png", image)
    if not ok:
        return b""
    return buf.tobytes()


# ----------------------------------------------------------------------
# 增量 PDF 写入器（每页处理完立即落盘，避免中途崩溃丢失全部结果）
# ----------------------------------------------------------------------
class IncrementalPdfWriter:
    """流式可搜索 PDF 写入器。

    每页处理完调用 append_page() 追加到内存文档，
    每 N 页（默认 5 页）才落盘一次，平衡内存释放频率与序列化开销。

    优化（基于日志实测）：
      - 旧实现每次 flush 都 save(garbage=4) 全部已有页，55 页时耗时 300s+
        （每隔 5 页出现一次异常高耗时，且越往后越慢）
      - 新实现改用 incremental=True 增量保存：只写入新增页，O(1) 开销
        不随总页数增长，且 fitz 增量保存会自动释放已写入页的图片缓存
      - 不再 close→open→save，避免重复序列化

    支持断点续传：若输出文件已存在（上次处理中断），打开已有 PDF 继续
    追加新页，跳过已处理页。existing_pages() 返回已有页数，供上层跳过。
    """

    def __init__(self, source_path: str, output_dir: str, render_dpi: int = 200,
                 flush_every: int = 5) -> None:
        import fitz

        self._fitz = fitz
        self.render_dpi = render_dpi
        self.scale = 72.0 / render_dpi
        self.out_path = _output_path(source_path, "pdf", output_dir)
        os.makedirs(os.path.dirname(self.out_path) or ".", exist_ok=True)
        # 断点续传：输出文件已存在则打开继续追加
        # 必须用 incremental=True 打开才能增量保存
        if os.path.isfile(self.out_path):
            try:
                self.doc = fitz.open(self.out_path)
                self._page_count = self.doc.page_count
                self._can_incremental = True  # 从文件打开，支持增量保存
                _log_writer(f"断点续传: 打开已有输出 {self.out_path}，已有 {self._page_count} 页")
            except Exception:
                _log_writer(f"已有输出文件损坏，重新创建: {self.out_path}")
                self.doc = fitz.open()
                self._page_count = 0
                self._can_incremental = False  # 新建空文档，首次 save 需全量
        else:
            self.doc = fitz.open()
            self._page_count = 0
            self._can_incremental = False  # 新建空文档，首次 save 需全量
        self._pages_since_flush = 0
        self._flush_every = max(1, flush_every)

    def existing_pages(self) -> int:
        """返回输出文件已有的页数（断点续传用）。"""
        return self._page_count

    def append_page(self, page: PageResult) -> None:
        """处理完一页后调用，将该页追加到输出 PDF。

        每 flush_every 页落盘一次，释放 fitz 内存中累积的图片缓存。
        增量保存（incremental=True）只写入新增页，开销 O(1) 不随总页数增长。
        """
        if page.image is None:
            return
        h, w = page.image.shape[:2]
        rect = self._fitz.Rect(0, 0, w * self.scale, h * self.scale)
        pdf_page = self.doc.new_page(width=rect.width, height=rect.height)
        # 叠加隐形文字层
        if not page.skipped:
            _insert_invisible_text(pdf_page, page, self.scale)
        # 插入原图
        img_bytes = _cv2_to_png_bytes(page.image)
        if img_bytes:
            pdf_page.insert_image(rect, stream=img_bytes)
        self._page_count += 1
        self._pages_since_flush += 1
        # 释放本页图片字节（已插入到 fitz doc，不再需要）
        del img_bytes
        # 每 flush_every 页落盘一次，释放内存
        if self._pages_since_flush >= self._flush_every:
            self._flush()
            self._pages_since_flush = 0

    def _flush(self) -> None:
        """增量保存到磁盘，释放 fitz 内存中的图片缓存。

        关键优化：用 incremental=True 增量保存，只写入新增页，
        避免 save(garbage=4) 每次重写全部已有页（旧实现 55 页时 300s+）。

        增量保存后 close→reopen 释放内存中的图片缓存：
          - fitz 在 insert_image 后会在内存中保留图片数据
          - close() 释放这些缓存，reopen 后只加载元数据
          - reopen 后必须标记 _can_incremental=True 才能继续增量保存
        """
        try:
            if self._can_incremental:
                # 增量保存：只写入新增页，O(1) 开销
                self.doc.save(self.out_path, incremental=True, encryption=0)
            else:
                # 首次保存（新建文档）：全量保存到文件
                self.doc.save(self.out_path, garbage=4, deflate=True)
                self._can_incremental = True
            # 关闭并重新打开，释放内存中的图片缓存
            self.doc.close()
            self.doc = self._fitz.open(self.out_path)
            # 从文件打开后支持增量保存
            self._can_incremental = True
        except PermissionError:
            raise RuntimeError(
                f"输出文件被占用，无法写入：\n  {self.out_path}\n\n"
                f"请关闭正在打开该文件的程序（如 PDF 阅读器、Word 等）后重试。"
            )
        except Exception as e:
            # 增量保存失败（如文件被移动/删除）：回退到全量保存到临时文件
            _log_writer(f"增量保存失败，回退全量保存: {e}")
            tmp_path = self.out_path + ".tmp"
            try:
                self.doc.save(tmp_path, garbage=4, deflate=True)
                self.doc.close()
                os.replace(tmp_path, self.out_path)
                self.doc = self._fitz.open(self.out_path)
                self._can_incremental = True
            except Exception:
                # 全量也失败：保留内存文档，下次再试
                self._safe_remove_tmp(tmp_path)
                _log_writer(f"全量保存也失败，保留内存文档继续: {e}")

    @staticmethod
    def _safe_remove_tmp(tmp_path: str) -> None:
        """安全删除临时文件，忽略失败。"""
        try:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass

    def close(self) -> str:
        """完成写入，返回输出路径。"""
        self._flush()
        self.doc.close()
        # 清理可能残留的临时文件
        self._safe_remove_tmp(self.out_path + ".tmp_a")
        self._safe_remove_tmp(self.out_path + ".tmp_b")
        return self.out_path

    @property
    def page_count(self) -> int:
        return self._page_count


# ----------------------------------------------------------------------
# DOCX 输出（替换段落文字保留样式 / 新建）
# ----------------------------------------------------------------------
def to_docx(result: DocumentResult) -> bytes:
    """生成 DOCX：原生文字段落 + OCR 识别文字 + 表格。

    表格处理：OCRResult.tables 中的 HTML 表格转换为 Word 原生表格，
    保证在 Word 中可编辑、可复制为表格结构。

    文本纠错：对每页的 OCR 行 text 单独做行级纠错（不合并行数，
    保持 line.coords 与表格 bbox 的空间对应关系正确）。

    去硬换行：纠错后按 _emit_page_lines_to_docx 把段内 OCR 物理行
    合并为逻辑段落（保留日期 / 序号 / 话题标签 / 短标题 / 空行等
    正常换行），每个段落作为一个 Word 段落输出，避免逐物理行换行。
    """
    from docx import Document
    from io import BytesIO
    import re

    doc = Document()
    doc.add_heading(os.path.basename(result.source_path), level=1)
    # 原生文字也做一次纠错
    native_text = result.native_text
    if native_text:
        native_text = _apply_fulltext_correction(
            native_text,
            source_name=f"{os.path.basename(result.source_path)} - DOCX原生"
        )
    if native_text:
        doc.add_heading("原生文字", level=2)
        for line in native_text.split("\n"):
            if line.strip():
                doc.add_paragraph(line)
    has_ocr = any(not p.skipped and p.ocr_result.lines for p in result.pages)
    if has_ocr:
        doc.add_heading("OCR 识别文字", level=2)
        corrector = _get_fulltext_corrector()
        for page in result.pages:
            if page.skipped:
                continue
            if len(result.pages) > 1:
                doc.add_heading(f"第 {page.page_no} 页", level=3)
            ocr = page.ocr_result
            # 先输出表格（按 bbox 顺序），再输出普通文字行
            # 表格内的文字行已在 PagePdfWriter 中被跳过，这里同样跳过
            table_bboxes = []
            if ocr.tables:
                for tbl in ocr.tables:
                    tbl_html = tbl.get("html", "")
                    tbl_bbox = tbl.get("bbox")
                    if tbl_html and tbl_bbox:
                        _add_html_table_to_docx(doc, tbl_html)
                        table_bboxes.append(tbl_bbox)
            # 先对本页所有 OCR 行的 text 做逐行纠错（merge_to_paragraphs=False 避免行数变化）
            page_line_texts = [ln.text for ln in ocr.lines]
            if corrector is not None and any(page_line_texts):
                try:
                    corrected = corrector.correct_lines(
                        page_line_texts, merge_to_paragraphs=False,
                    )
                    # 逐行写回（DOCX 场景保留物理行数，不做合并）
                    for idx, new_t in enumerate(corrected.lines):
                        if idx < len(ocr.lines):
                            ocr.lines[idx].text = new_t
                except Exception as e:
                    logger.warning("DOCX页级文本纠错异常（跳过）: %s", e)
            # 收集非表格区域的文字行（跳过表格区域内已被表格呈现的行）
            page_text_lines: List[str] = []
            for line in ocr.lines:
                if not line.text:
                    continue
                # 跳过表格区域内的行（表格已用 Word 表格呈现）
                if table_bboxes and line.coords:
                    cx = sum(p[0] for p in line.coords) / len(line.coords)
                    cy = sum(p[1] for p in line.coords) / len(line.coords)
                    in_table = any(
                        tx0 <= cx <= tx1 and ty0 <= cy <= ty1
                        for tx0, ty0, tx1, ty1 in table_bboxes
                    )
                    if in_table:
                        continue
                page_text_lines.append(line.text)
            # 合并硬换行为段落，输出到 docx（保留 Markdown 表格块结构）
            _emit_page_lines_to_docx(doc, page_text_lines)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_html_table_to_docx(doc, html: str) -> None:
    """把 HTML 表格转换为 Word 原生表格。

    支持 SLANet 输出的 HTML 表格结构，包括合并单元格（colspan/rowspan）。
    """
    import re
    rows = re.findall(r"<tr[^>]*>(.*?)</tr>", html, re.DOTALL | re.IGNORECASE)
    if not rows:
        return
    # 解析所有行，收集单元格信息
    parsed_rows = []
    max_cols = 0
    for row in rows:
        cells = re.findall(
            r"<t[hd][^>]*>(.*?)</t[hd]>", row, re.DOTALL | re.IGNORECASE
        )
        if not any(re.sub(r"<[^>]+>", "", c).strip() for c in cells):
            continue
        parsed_rows.append(cells)
        max_cols = max(max_cols, len(cells))
    if not parsed_rows or max_cols == 0:
        return
    # 创建 Word 表格
    table = doc.add_table(rows=len(parsed_rows), cols=max_cols)
    table.style = "Table Grid"
    for ri, cells in enumerate(parsed_rows):
        for ci, cell_html in enumerate(cells):
            if ci >= max_cols:
                break
            text = re.sub(r"<[^>]+>", "", cell_html).strip()
            table.cell(ri, ci).text = text


def _add_markdown_table_to_docx(doc, md_text: str) -> None:
    """把 Markdown 表格文本转换为 Word 原生表格。

    Markdown 表格格式：
    | 表头1 | 表头2 |
    | --- | --- |
    | 单元格1 | 单元格2 |
    """
    lines = md_text.strip().split("\n")
    if len(lines) < 2:
        doc.add_paragraph(md_text)
        return
    # 解析行
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        # 跳过分隔行 | --- | --- |
        if re.match(r"^\|[\s\-:|]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    if not rows:
        doc.add_paragraph(md_text)
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for ri, cells in enumerate(rows):
        for ci, text in enumerate(cells):
            if ci < max_cols:
                table.cell(ri, ci).text = text


def _emit_page_lines_to_docx(doc, text_lines: List[str]) -> None:
    """把一页文字行输出到 docx，合并段内硬换行，保留 Markdown 表格块。

    - 连续以 ``|`` 开头的行视为一个 Markdown 表格块，整块传给
      _add_markdown_table_to_docx 还原为 Word 表格（避免表格行被合并）。
    - 其余文本行先按 merge_lines_to_paragraphs 合并为段落（去硬换行），
      再逐段添加为 Word 段落。
    """
    i = 0
    n = len(text_lines)
    while i < n:
        if text_lines[i].strip().startswith("|"):
            # 收集连续的 Markdown 表格行
            block: List[str] = []
            while i < n and text_lines[i].strip().startswith("|"):
                block.append(text_lines[i])
                i += 1
            _add_markdown_table_to_docx(doc, "\n".join(block))
            continue
        # 收集普通文本段（到下一个表格块或末尾）
        seg: List[str] = []
        while i < n and not text_lines[i].strip().startswith("|"):
            seg.append(text_lines[i])
            i += 1
        for para in merge_lines_to_paragraphs(seg):
            if para.strip():
                doc.add_paragraph(para)


# ----------------------------------------------------------------------
# 统一保存入口
# ----------------------------------------------------------------------
def save_output(
    result: DocumentResult,
    fmt: str,
    output_dir: str,
    render_dpi: int = 200,
) -> str:
    """按指定格式保存输出，返回输出文件路径。

    fmt: original | txt | markdown | json | searchable_pdf | docx
    """
    src = result.source_path
    src_ext = os.path.splitext(src)[1].lower()

    if fmt == "original":
        # 还原为原格式
        if src_ext == ".pdf":
            return _save_bytes(src, "pdf", output_dir, to_searchable_pdf(result, render_dpi))
        if src_ext == ".docx":
            return _save_bytes(src, "docx", output_dir, to_docx(result))
        # 图片 -> txt
        return _save_text(src, "txt", output_dir, to_txt(result))

    if fmt == "txt":
        return _save_text(src, "txt", output_dir, to_txt(result))
    if fmt == "markdown":
        return _save_text(src, "md", output_dir, to_markdown(result))
    if fmt == "json":
        return _save_text(src, "json", output_dir, to_json(result))
    if fmt == "searchable_pdf":
        return _save_bytes(src, "pdf", output_dir, to_searchable_pdf(result, render_dpi))
    if fmt == "docx":
        return _save_bytes(src, "docx", output_dir, to_docx(result))

    raise ValueError(f"不支持的输出格式: {fmt}")


def _save_text(source_path: str, ext: str, output_dir: str, content: str) -> str:
    path = _output_path(source_path, ext, output_dir)
    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    return path


def _save_bytes(source_path: str, ext: str, output_dir: str, data: bytes) -> str:
    path = _output_path(source_path, ext, output_dir)
    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    with open(path, "wb") as f:
        f.write(data)
    return path
