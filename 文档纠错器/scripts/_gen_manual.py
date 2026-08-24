# -*- coding: utf-8 -*-
"""生成文档纠错器用户使用手册（docx）。

保姆级、面向无计算机基础用户、无废话。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "test_output")
os.makedirs(OUT_DIR, exist_ok=True)
OUT_PATH = os.path.join(OUT_DIR, "_用户使用手册.docx")


def set_font(run, name="微软雅黑", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 中文字体
    from docx.oxml.ns import qn
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=22, bold=True)
    p.space_after = Pt(6)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=11, color=(128, 128, 128))


def h1(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=16, bold=True, color=(0, 51, 102))


def h2(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=13, bold=True, color=(0, 80, 160))


def h3(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r, size=11, bold=True)


def body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.space_after = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    r = p.add_run(text)
    set_font(r, size=10.5)


def bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.75)
    p.space_after = Pt(2)
    r = p.add_run(("• 「 if level == 0 else 「◦ ") + text)
    set_font(r, size=10.5)


def numbered(doc, n, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.space_after = Pt(2)
    r = p.add_run(f"{n}. {text}")
    set_font(r, size=10.5)


def note(doc, text):
    """提示框：浅灰背景"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.space_before = Pt(4)
    p.space_after = Pt(4)
    r = p.add_run("【提示】" + text)
    set_font(r, size=10, color=(100, 100, 100))


def warning(doc, text):
    """警告框"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.space_before = Pt(4)
    p.space_after = Pt(4)
    r = p.add_run("【注意】" + text)
    set_font(r, size=10, color=(192, 0, 0))


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_font(r, size=10, bold=True)
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            set_font(r, size=10)
    doc.add_paragraph()  # 表后空行


def build():
    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 封面
    add_title(doc, "文档纠错器")
    add_subtitle(doc, "用户使用手册")
    doc.add_paragraph()
    add_subtitle(doc, "基于 DeepSeek V4 大模型的中文文档智能校对工具")
    doc.add_paragraph()
    add_subtitle(doc, "适用于 Windows 10 / 11")
    doc.add_paragraph()
    doc.add_paragraph()
    body(doc, "本手册面向无计算机使用经验的用户，按顺序阅读即可上手。")

    doc.add_page_break()

    # 目录提示
    h1(doc, "目录")
    for item in [
        "一、什么是文档纠错器",
        "二、安装与首次配置",
        "三、准备 DeepSeek API 密钥",
        "四、准备讯飞 OCR 密钥（可选）",
        "五、界面认识",
        "六、校对一篇文档（完整流程）",
        "七、OCR 图片识别纠错模式",
        "八、跳转页与翻页提示",
        "九、错误处理：采纳、忽略、撤回",
        "十、导出修正后的文档",
        "十一、进度保存与恢复",
        "十二、模型选择与费用估算",
        "十三、常见问题",
    ]:
        body(doc, item)

    doc.add_page_break()

    # 一、什么是文档纠错器
    h1(doc, "一、什么是文档纠错器")
    body(doc, "文档纠错器是一款帮助您检查中文文档错误的桌面软件。它能自动找出文档中的：")
    bullet(doc, "错别字（如「部署「写成「布署「）")
    bullet(doc, "语法语病（如成分残缺、搭配不当）")
    bullet(doc, "标点误用（如引号方向错误、顿号逗号混用）")
    bullet(doc, "数据与逻辑矛盾（如前后年份不一致、数字对不上）")
    body(doc, "找到错误后，您可以选择采纳或忽略，最后导出修正后的文档，原格式（加粗、表格等）保持不变。")
    note(doc, "本软件需要联网调用 AI 大模型，请确保电脑能正常上网。")

    # 二、安装与首次配置
    h1(doc, "二、安装与首次配置")
    h2(doc, "2.1 获取软件")
    body(doc, "您会收到一个名为「文档纠错器.exe「的文件。这就是完整的软件，无需安装，双击即可运行。")
    warning(doc, "请勿删除或重命名该文件所在文件夹中的 config.json 文件——那是您的配置和密钥。")

    h2(doc, "2.2 首次启动")
    numbered(doc, 1, "双击「文档纠错器.exe「。")
    numbered(doc, 2, "如果弹出「Windows 已保护你的电脑「提示，点击「更多信息「，再点击「仍要运行「。")
    numbered(doc, 3, "软件打开后，界面是空白的——这是正常的，因为还没有打开文档。")
    numbered(doc, 4, "此时请先配置密钥（见下一节），否则无法使用校对功能。")

    h2(doc, "2.3 配置密钥（必须）")
    body(doc, "软件需要两把「钥匙「才能工作：")
    bullet(doc, "DeepSeek API 密钥——用于 AI 校对（必须）")
    bullet(doc, "讯飞 OCR 密钥——用于识别图片和扫描版 PDF（可选，不识别图片则不需要）")
    body(doc, "配置方法：")
    numbered(doc, 1, "点击菜单栏的「设置「。")
    numbered(doc, 2, "在「DeepSeek「标签页填入 API Key（一串以 sk- 开头的字符）。")
    numbered(doc, 3, "如需 OCR 功能，切换到「讯飞 OCR「标签页填入 AppID、API Key、API Secret。")
    numbered(doc, 4, "点击「保存「。")

    # 三、DeepSeek 密钥
    h1(doc, "三、准备 DeepSeek API 密钥")
    h2(doc, "3.1 注册账号")
    numbered(doc, 1, "打开浏览器，访问 https://platform.deepseek.com")
    numbered(doc, 2, "注册一个账号（手机号即可）。")
    numbered(doc, 3, "登录后，在左侧菜单找到「API Keys「。")
    h2(doc, "3.2 创建密钥")
    numbered(doc, 1, "点击「创建 API Key「。")
    numbered(doc, 2, "复制生成的那串字符（以 sk- 开头）。")
    numbered(doc, 3, "粘贴到软件的「设置 → DeepSeek → API Key「输入框里。")
    h2(doc, "3.3 充值")
    body(doc, "DeepSeek 采用按量付费，需要先充值才能使用。")
    numbered(doc, 1, "在平台左侧菜单找到「费用「。")
    numbered(doc, 2, "点击「充值「，建议先充 10 元试用（足够校对几万字）。")
    note(doc, "校对 1 万字大约花费 5 分钱（flash 模型），费用很低。详见第十二章。")

    # 四、讯飞 OCR
    h1(doc, "四、准备讯飞 OCR 密钥（可选）")
    body(doc, "只有当您需要校对图片或扫描版 PDF 时才需要配置。普通的 Word 文档、TXT 文本不需要。")
    h2(doc, "4.1 注册")
    numbered(doc, 1, "访问 https://www.xfyun.cn")
    numbered(doc, 2, "注册账号并完成实名认证。")
    h2(doc, "4.2 创建应用")
    numbered(doc, 1, "进入控制台，找到「印刷文字识别「服务。")
    numbered(doc, 2, "领取免费额度（新用户通常有免费体验量）。")
    numbered(doc, 3, "创建一个应用，获取三个信息：AppID、API Key、API Secret。")
    h2(doc, "4.3 填入软件")
    body(doc, "将上述三个信息填入软件「设置 → 讯飞 OCR「标签页，保存即可。")

    doc.add_page_break()

    # 五、界面认识
    h1(doc, "五、界面认识")
    body(doc, "软件界面从上到下、从左到右分为以下几个区域：")
    h2(doc, "5.1 顶部工具栏")
    bullet(doc, "「打开」：选择要校对的文档")
    bullet(doc, "「开始纠错」：对当前页启动 AI 校对")
    bullet(doc, "「图片识别纠错」：切换到 OCR 专用纠错模式（按钮按下=启用）")
    bullet(doc, "「下一页」：提交当前页修改，进入下一页")
    bullet(doc, "「跳转页」：直接跳到指定页码开始校对")
    bullet(doc, "「导出」：保存修正后的文档")
    bullet(doc, "「设置」：配置密钥、模型、规则")

    h2(doc, "5.2 左侧——文档预览区")
    body(doc, "显示当前文档内容。校对出错误后，对应文字会高亮显示。点击右侧错误列表中的某条，"
              "左侧会自动滚动到对应位置。")

    h2(doc, "5.3 右侧——错误面板")
    body(doc, "分为「未确认「和「已确认「两个标签页：")
    bullet(doc, "未确认：AI 找到的错误列表，等待您处理")
    bullet(doc, "已确认：您已采纳的修改，可以在这里撤回")
    body(doc, "每条错误显示：页码、类型、置信度、原文、建议、理由。文字太长会显示省略号，"
              "双击某条可查看完整内容。")

    h2(doc, "5.4 底部——状态栏")
    body(doc, "显示当前页码、正在分析的块号、已等待时间、累计消耗的 token 数量。")

    # 六、完整流程
    doc.add_page_break()
    h1(doc, "六、校对一篇文档（完整流程）")
    body(doc, "以下是从头到尾校对一篇文档的完整步骤：")

    h2(doc, "步骤 1：打开文档")
    numbered(doc, 1, "点击工具栏「打开」按钮。")
    numbered(doc, 2, "在弹出的窗口中找到您的文档（默认显示「我的文档「文件夹）。")
    numbered(doc, 3, "选中文档，点击「打开「。")
    numbered(doc, 4, "左侧预览区会显示文档内容。")
    note(doc, "支持的格式：Word 文档（.docx）、纯文本（.txt）、文字版 PDF。"
              "图片和扫描版 PDF 需要 OCR 功能。")

    h2(doc, "步骤 2：开始纠错")
    numbered(doc, 1, "点击工具栏「开始纠错」按钮。")
    numbered(doc, 2, "底部状态栏会显示「正在分析第 1 块…已等待 Xs「，表示 AI 正在工作。")
    numbered(doc, 3, "等待几秒到几十秒（取决于文档长度和网络速度）。")
    numbered(doc, 4, "完成后，右侧「未确认「标签页会列出找到的错误。")

    h2(doc, "步骤 3：查看与处理错误")
    body(doc, "对每条错误，您可以：")
    bullet(doc, "「采纳」：接受这条修改建议，原文会被替换")
    bullet(doc, "「忽略」：不采纳，跳过这条")
    bullet(doc, "「一键替换」：把当前筛选结果全部采纳（适合批量处理）")
    bullet(doc, "双击某条：查看完整的原文、建议和理由，可复制")
    note(doc, "点击某条错误后，左侧预览会自动滚动到对应位置并高亮。")

    h2(doc, "步骤 4：翻到下一页")
    numbered(doc, 1, "当前页的错误处理完后，点击「下一页」。")
    numbered(doc, 2, "软件会提交当前页的修改，然后自动开始下一页的校对。")
    warning(doc, "如果当前页还有未确认的错误，翻页前会弹窗提醒。"
                 "翻页后未处理的错误将被放弃（不会修改原文）。")

    h2(doc, "步骤 5：导出文档")
    numbered(doc, 1, "全部页面处理完后（或您想中途保存时），点击「导出」。")
    numbered(doc, 2, "默认保存到「我的文档/DocProof/导出/「文件夹。")
    numbered(doc, 3, "文件名会自动加上「_已修正「后缀，不会覆盖原文档。")
    numbered(doc, 4, "导出后进度文件自动清理。")

    # 七、OCR 模式
    doc.add_page_break()
    h1(doc, "七、OCR 图片识别纠错模式")
    body(doc, "如果您的文档是通过扫描、拍照后用 OCR 软件识别得到的，文字中可能存在一些"
              "OCR 特有的错误（如形近字误识、数字识别错误、漏字漏行、重复识别、乱码等）。")
    h2(doc, "如何使用")
    numbered(doc, 1, "打开文档后，点击工具栏「图片识别纠错」按钮（按钮被按下表示已启用）。")
    numbered(doc, 2, "点击「开始纠错」，AI 会使用 OCR 专用提示词进行校对。")
    numbered(doc, 3, "处理完成后，再次点击该按钮可切换回常规模式。")
    h2(doc, "OCR 模式能检出的错误类型")
    add_table(doc, ["错误类型", "举例", "检出能力"], [
        ["形近字误识", "末/未、己/已、戌/戊", "强"],
        ["数字误识", "3→8、1948→1949", "中等（逻辑矛盾能检出，纯数字难判）"],
        ["标点误识", "句号识为逗号、引号方向错", "中等"],
        ["漏字/漏行/漏句", "城门名漏识、数量词漏识", "强（单字漏识），弱（整句漏识）"],
        ["重复识别", "同一行被识别两次", "强"],
        ["乱码字符", "无意义符号替代汉字", "强"],
        ["残句拼接", "句子只识别了一半", "强"],
        ["两栏穿插/段落混乱", "版面结构错乱", "弱（纯文本无法感知版面）"],
    ])
    note(doc, "OCR 模式对字词级错误检出效果好，对版面结构级错误（如两栏文字穿插）效果有限。"
              "版面错乱问题需要从源头用更好的 OCR 软件重新识别。")

    # 八、跳转页
    h1(doc, "八、跳转页与翻页提示")
    h2(doc, "8.1 跳转页")
    body(doc, "如果您不需要从头校对，想直接从某一页开始：")
    numbered(doc, 1, "点击工具栏「跳转页」按钮。")
    numbered(doc, 2, "输入目标页码（如「3「），点击确定。")
    numbered(doc, 3, "软件会提交当前页修改，然后直接跳到第 3 页开始校对。")
    body(doc, "支持跳过中间页（如从第 1 页跳到第 7 页），也支持回跳到已经处理过的页。")

    h2(doc, "8.2 翻页提示")
    body(doc, "当当前页还有未确认的错误时，您点击「下一页」或「跳转页」，会弹出提示：")
    body(doc, "「第 X 页还有 N 条未确认错误。翻页后将放弃这些错误（不修改原文），是否继续？「")
    body(doc, "选择「是「则翻页（未处理的错误被放弃），选择「否「则留在当前页继续处理。")

    # 九、错误处理
    doc.add_page_break()
    h1(doc, "九、错误处理：采纳、忽略、撤回")
    h2(doc, "9.1 筛选错误")
    body(doc, "右侧「未确认「面板顶部有筛选下拉框：")
    bullet(doc, "按类型：全部 / 错别字 / 语法 / 标点 / 逻辑")
    bullet(doc, "按置信度：全部 / 明确 / 存疑")
    note(doc, "「明确「= AI 很确定这是错误；「存疑「= AI 不太确定，建议人工重点判断。"
              "存疑行整行显示为灰色。")

    h2(doc, "9.2 单条处理")
    bullet(doc, "「采纳」：接受建议，原文被替换为建议内容")
    bullet(doc, "「忽略」：不采纳，该条从列表消失")
    bullet(doc, "双击：弹出详情对话框，可复制建议或理由")

    h2(doc, "9.3 批量处理")
    bullet(doc, "「一键替换」：把当前筛选结果全部采纳")
    note(doc, "例如先筛选「错别字「，再点「一键替换「，可一次性采纳所有错别字修改。")

    h2(doc, "9.4 撤回")
    body(doc, "切换到「已确认「标签页：")
    bullet(doc, "单条撤回：选中某条，点击「撤回「")
    bullet(doc, "按批次回滚：点击「回滚此批次「，撤销同一次「一键替换「的所有修改")

    # 十、导出
    h1(doc, "十、导出修正后的文档")
    h2(doc, "10.1 导出方式")
    numbered(doc, 1, "点击工具栏「导出」按钮。")
    numbered(doc, 2, "默认保存到「我的文档/DocProof/导出/「文件夹（自动创建）。")
    numbered(doc, 3, "文件名自动加「_已修正「后缀，不会覆盖原文件。")
    numbered(doc, 4, "如需保存到其他位置，在保存窗口中自行选择。")
    h2(doc, "10.2 格式保留")
    add_table(doc, ["原格式", "导出格式", "格式保留情况"], [
        ["Word（.docx）", "Word（.docx）", "保留加粗、斜体、表格等格式"],
        ["纯文本（.txt）", "纯文本（.txt）", "纯文本无格式"],
        ["PDF / 图片", "Word（.docx）", "转为 Word 格式"],
    ])
    warning(doc, "导出后进度文件会自动清理。如需继续校对，需要重新打开文档。")

    # 十一、进度
    h1(doc, "十一、进度保存与恢复")
    body(doc, "软件会自动保存您的校对进度，无需手动操作。")
    h2(doc, "11.1 自动保存")
    bullet(doc, "每翻一页、每次纠错完成都会自动保存进度。")
    bullet(doc, "进度文件在源文档同目录下，名为「_stem.docproof.json「。")
    h2(doc, "11.2 恢复进度")
    body(doc, "如果软件中途崩溃或被误关：")
    numbered(doc, 1, "重新打开软件，打开同一个文档。")
    numbered(doc, 2, "会弹窗提示「是否恢复上次的进度「，显示已修改条数和保存时间。")
    numbered(doc, 3, "点击「是「，自动跳转到上次处理到的页码。")
    h2(doc, "11.3 进度失效")
    body(doc, "以下情况进度会被判为失效，不恢复：")
    bullet(doc, "源文档在软件外被修改过（如用 Word 编辑后保存）")
    bullet(doc, "完成导出后（进度文件自动删除）")

    # 十二、费用
    doc.add_page_break()
    h1(doc, "十二、模型选择与费用估算")
    h2(doc, "12.1 两个模型")
    add_table(doc, ["模型", "特点", "适合场景", "推荐"], [
        ["flash（默认）", "非思考模式，快、省", "日常校对", "首选"],
        ["pro", "思考模式，更细但误报多", "成书定稿前最后一遍", "需人工复核存疑项"],
    ])
    note(doc, "在「设置 → DeepSeek → 模型「中切换。建议日常用 flash，定稿前用 pro 跑一遍。")

    h2(doc, "12.2 每万字费用")
    body(doc, "按 DeepSeek 官方定价（2026年7月，1美元≈7.2元）：")
    add_table(doc, ["场景", "flash", "pro"], [
        ["平峰期（最常见）", "约 ¥0.054", "约 ¥0.247"],
        ["高峰期", "约 ¥0.108", "约 ¥0.494"],
    ])

    h2(doc, "12.3 高峰时段")
    body(doc, "DeepSeek 采用峰谷分时计费：")
    bullet(doc, "高峰时段：每天 09:00–12:00 和 14:00–18:00（价格为平峰的 2 倍）")
    bullet(doc, "平峰时段：其余时间（含夜间和午休）")
    note(doc, "大批量校对建议安排在 18:00 后或 09:00 前，费用减半。单篇文档随时跑即可。")

    h2(doc, "12.4 常见文档费用估算（flash 平峰）")
    add_table(doc, ["文档类型", "字数", "费用"], [
        ["单篇公文", "3,000 字", "约 ¥0.016"],
        ["一卷志稿", "30,000 字", "约 ¥0.162"],
        ["整部县志（10卷）", "300,000 字", "约 ¥1.62"],
    ])
    note(doc, "开启二次复核约增加 30% 费用。上述估算仅供参考。")

    h2(doc, "12.3 推荐配置方案")
    h3(doc, "日常使用（推荐）")
    bullet(doc, "模型：flash")
    bullet(doc, "二次复核：开启")
    bullet(doc, "大批量任务安排在夜间")
    h3(doc, "严格审稿（定稿前）")
    bullet(doc, "模型：pro")
    bullet(doc, "二次复核：开启")
    bullet(doc, "务必人工复核所有「存疑「项")
    h3(doc, "成本敏感（个人用户）")
    bullet(doc, "模型：flash")
    bullet(doc, "二次复核：关闭")
    bullet(doc, "仅在夜间平峰期使用")

    # 十三、常见问题
    doc.add_page_break()
    h1(doc, "十三、常见问题")

    h2(doc, "Q1：双击 exe 没反应？")
    body(doc, "请等待 10-30 秒。软件首次启动需要解压运行库，稍慢。如果超过 1 分钟仍无反应，"
              "请检查是否有杀毒软件拦截。")

    h2(doc, "Q2：提示「API 密钥错误「？")
    body(doc, "请检查「设置 → DeepSeek → API Key「是否正确（以 sk- 开头），"
              "以及账户是否有余额。")

    h2(doc, "Q3：提示「网络连接失败「？")
    body(doc, "请检查电脑是否能正常上网。如果使用公司内网，可能需要配置代理。")

    h2(doc, "Q4：校对速度很慢？")
    body(doc, "速度取决于文档长度和网络状况。一般 1 页（约 1000 字）需要 10-30 秒。"
              "如果超过 1 分钟，可能是网络问题或 DeepSeek 服务器繁忙，请稍后重试。")

    h2(doc, "Q5：AI 找出的错误不准？")
    body(doc, "AI 校对不可能 100% 准确。建议：")
    bullet(doc, "优先看「明确「级别的错误")
    bullet(doc, "「存疑「级别的错误请人工判断")
    bullet(doc, "开启二次复核可降低误报")
    bullet(doc, "flash 模型误报较少，pro 模型更细但误报更多")

    h2(doc, "Q6：能校对英文文档吗？")
    body(doc, "本软件针对中文文档优化，英文校对效果不保证。")

    h2(doc, "Q7：打开 PDF 提示需要 OCR？")
    body(doc, "说明该 PDF 是扫描版（图片形式），需要配置讯飞 OCR 密钥。"
              "文字版 PDF（可直接复制文字的）不需要 OCR。")

    h2(doc, "Q8：导出的文档格式变了？")
    body(doc, "Word 文档导出后保留原格式（加粗、表格等）。PDF 和图片导出为 Word 格式。"
              "如果发现格式异常，请检查原文档是否包含复杂排版。")

    h2(doc, "Q9：进度文件能删吗？")
    body(doc, "可以。「_stem.docproof.json「是进度文件，删除后只是无法恢复进度，"
              "不影响文档本身。导出后会自动删除。")

    h2(doc, "Q10：config.json 能删吗？")
    warning(doc, "不建议删除。config.json 保存着您的密钥和设置，删除后需要重新配置。"
                 "如果只是想重置设置，可以在软件「设置「里修改后保存。")

    # 结尾
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— 本手册结束 —")
    set_font(r, size=10, color=(128, 128, 128))

    doc.save(OUT_PATH)
    print(f"用户手册已生成：{OUT_PATH}")
    print(f"文件大小：{os.path.getsize(OUT_PATH) / 1024:.1f} KB")


if __name__ == "__main__":
    build()
