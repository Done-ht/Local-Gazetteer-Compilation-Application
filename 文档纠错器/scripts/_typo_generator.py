# -*- coding: utf-8 -*-
"""生成含隐蔽错别字的测试文档，并测试 flash/pro 的错别字校对能力。

错别字埋设策略（由隐蔽到明显）：
1. **同音字**：如"部署"→"布署"、"迫不及待"→"迫不急待"、"冒犯"→"冒反"
2. **形近字**：如"己/已/巳"、"戊/戌/戍"、"玊/玉"、"姬/姫"
3. **成语错字**：如"贻笑大方"→"贻笑大房"、"罄竹难书"→"磬竹难书"
4. **专有名词错字**：如"邯郸"→"甘郸"、"嘉峪关"→"嘉峪关"（这种很难）
5. **同义字混淆**：如"登录"→"登陆"、"账目"→"帐目"

每个测试文档约 600-1000 字，包含 3-5 处错别字，明确记录埋设位置便于评估。

输出：
    test_output/_typos/*.docx           — 含错别字的测试文档
    test_output/_typos/_typo_seeds.json — 错别字埋设清单（ground truth）
"""
import json
import os
import sys
import time

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_ROOT)
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt

OUT_DIR = os.path.join(PROJECT_ROOT, "test_output", "_typos")
os.makedirs(OUT_DIR, exist_ok=True)

# ===== 测试文档与错别字埋设 =====
# 每篇文档：dict(doc_name, title, paragraphs[list[str]], typos[list[(bad, good, hint)]])
# typos 中的 (bad, good) 必须在 paragraphs 中实际出现 bad，便于评估

TEST_DOCS = [
    {
        "doc_name": "typos_01_homophone.docx",
        "title": "测试文档 1：同音字错别字（公文类）",
        "desc": "埋设 5 处同音字错别字",
        "paragraphs": [
            "关于推进乡村振兴重点工作的实施意见（节录）",
            "为深入贯彻落实党中央、国务院关于乡村振兴的战略布署，结合我县实际，现就2025年重点工作提出如下意见。",
            "一、加强基础设施建设。各乡镇要按照县政府的统一布署，加快推进农村道路、供水、供电等基础设施建设，确保年底前完成全部既定目标。",
            "二、强化产业扶持。重点支持特色农产品种植与加工，对带动作用强的龙头企业给予财政贴息，对亟待转型升级的中小企业给予政策倾斜，迫不急待地推动一二三产业融合发展。",
            "三、加强人才队伍建设。鼓励大学生回乡创业，对返乡创业人员给予创业补贴和税收优惠；建立县乡村三级人才交流机制，让技术人材下沉到基层。",
            "四、严格考核问责。各乡镇、各部门要切实扛起责任，对工作不力、进展缓慢的，将严肃追责问责，决不徇私舞弊，确保政令畅通、令行禁止。",
            "五、附则。本意见自发布之日起施行，由县乡村振兴局负责解释。各乡镇、各部门要根据本意见制定具体实施方案，并报县乡村振兴局备案。",
        ],
        "typos": [
            ("布署", "部署", "同音错别字：应为'部署'，'布署'是非规范写法"),
            ("布署", "部署", "同音错别字：第二次出现"),
            ("迫不急待", "迫不及待", "成语错别字：'迫不及待'误作'迫不急待'"),
            ("人材", "人才", "同音错别字：'人才'误作'人材'（古语可通用，现代公文应作'人才'）"),
            ("徇私舞弊", "徇私舞弊", "（对照项：此处正确，不应报错）"),
        ],
    },
    {
        "doc_name": "typos_02_form_similar.docx",
        "title": "测试文档 2：形近字错别字（史志类）",
        "desc": "埋设 5 处形近字错别字",
        "paragraphs": [
            "《某县志·卷七·人物》（节录）",
            "张某，字怀璞，世居本县长春乡。幼颖悟，年十二即通《论语》《孟子》，乡里称奇。弱冠补博士弟子员，旋食饩。",
            "清道光二十七年（1847年）进士，授翰林院编修。在职期间，先后典试湖南、湖北，所至拔擢多佳士，朝野仰之。",
            "咸丰初，擢侍讲学士，转侍读学士。时太平天国起事，江南骚动，张公上疏陈防守机宜，言甚切直，权臣忌之。",
            "公为人清介，居家孝友。尝捐资修治家庙，赡养族人；置义田以济贫乏，乡党义之。著《长春斋集》二十卷行世。",
            "同治五年卒于家，年六十九。门人私谥曰文端公，祀乡贤祠。其墓在县西十五里白羊岗，松柏森然，过者式焉。",
            "赞曰：张公以文学起家，致身清要，而能固穷守道，不谄权贵，可谓浊世之君子矣。其言行事迹，载本志《人物传》。",
        ],
        "typos": [
            ("怀璞", "怀璞", "（对照项：正确，不应报）"),
            ("食饩", "食饩", "（对照项：生僻字但正确，不应报）"),
            ("道光二十七年（1847年）", "道光二十七年（1847年）", "（对照项：实际正确，不应报）"),
            ("侍讲学士", "侍讲学士", "（对照项：正确，不应报）"),
            ("白羊岗", "白羊岗", "（对照项：正确，不应报）"),
        ],
        # 注意：本文档故意全部正确，测试模型是否会"无中生有"误报
        "typos_actual": 0,
    },
    {
        "doc_name": "typos_03_idiom.docx",
        "title": "测试文档 3：成语错别字（评论类）",
        "desc": "埋设 6 处成语错别字",
        "paragraphs": [
            "在新的历史起点上推进县域治理现代化（评论）",
            "近年来，我县在经济社会发展方面取得了显著成绩，但也要清醒看到，发展不平衡不充分的问题仍然突出，需要我们破斧沉舟、攻坚克难。",
            "首先，要旗帜鲜明地坚持以人民为中心的发展思想。各级干部不能因为取得一点成绩就沾沾自喜，更不能因噎废食、固步自封，要始终保持昂扬向上的精神状态。",
            "其次，要深化改革开放。改革进入深水区，剩下的都是难啃的硬骨头。我们要以破釜沉舟的勇气、壮士断腕的决心，敢于向积弊开刀，敢于向痼疾动刀。",
            "再次，要统筹发展和安全。安全是发展的前提，发展是安全的保障。各级各部门要未雨绸缪，建立健全风险预警机制，把矛盾化解在萌芽状态。",
            "最后，要加强党的全面领导。党是领导一切的。各级党组织要切实扛起主责主业，发挥战斗堡垒作用，让党员的先锋模范作用惠及每个角落。",
            "总之，县域治理现代化是一项系统工程，需要我们以钉钉子精神，一张蓝图绘到底，才能在中华民族伟大复兴的征程中谱写出无愧于时代的篇章。",
        ],
        "typos": [
            ("破斧沉舟", "破釜沉舟", "成语错别字：'破釜沉舟'误作'破斧沉舟'（釜：锅；不是斧头）"),
            ("沾沾自喜", "沾沾自喜", "（对照项：正确）"),
            ("因噎废食", "因噎废食", "（对照项：正确）"),
            ("固步自封", "故步自封", "形近错别字：'故步自封'规范写法应作'故步'（'固步'是常见变体，规范推荐'故步'）"),
            ("破釜沉舟", "破釜沉舟", "（对照项：此处正确，第二次出现）"),
            ("未雨绸缪", "未雨绸缪", "（对照项：正确）"),
            ("惠及", "惠及", "（对照项：正确）"),
            ("谱写出", "谱写出", "（对照项：正确）"),
        ],
        "typos_actual": 2,  # "破斧沉舟" + "固步自封"
    },
    {
        "doc_name": "typos_04_mixed.docx",
        "title": "测试文档 4：混合错别字（综合）",
        "desc": "埋设 7 处各类错别字（含同音、形近、成语、专名）",
        "paragraphs": [
            "【县情介绍】",
            "本县位于皖南山区，地处黄山、九华山之间，历史悠久，文化底蕴深厚。全县总面积1268平方公里，辖9镇3乡，人口约23万。",
            "县境地势南高北低，最高峰清凉峰海拔1787米。主要河流有青弋江、水阳江，皆汇入长江。气候属亚热带季风气候，四季分明，雨量充沛。",
            "本县自唐代建置以来，已历1300余年。历代名人辈出，宋代有理学家朱熹在此讲学，明代有医学家汪机世居本县，清代有学者戴震埋首著述。",
            "近年来，本县立足生态优势，大力发展乡村旅游。黄山脚下、太平湖畔，民宿星罗棋布，年接待游客超过200万人次，旅游综合收入近15亿元。",
            "县内交通便捷，京福高速、商合杭高铁穿境而过，至合肥、杭州均不超过2小时车程。教育医疗资源完善，现有中学12所、小学35所、县级医院2所。",
            "欢迎各界朋友来我县投资兴业、旅游观光，共谋发展、同谱华章。我们将竭诚为您提供优质服务，让您的投资得到丰厚回报、让您的旅程留连忘返。",
        ],
        "typos": [
            ("皖南山区", "皖南山区", "（对照项：正确）"),
            ("青弋江", "青弋江", "（对照项：正确，注意'弋'易与'戈'混）"),
            ("埋首著述", "埋首著述", "（对照项：正确）"),
            ("星罗棋布", "星罗棋布", "（对照项：正确）"),
            ("留连忘返", "流连忘返", "同音错别字：'流连忘返'规范写法应作'流连'（'留连'为异形词，规范推荐'流连'）"),
        ],
        "typos_actual": 1,  # "留连忘返"
    },
    {
        "doc_name": "typos_05_subtle.docx",
        "title": "测试文档 5：极隐蔽错别字（高频混淆字）",
        "desc": "埋设 6 处极隐蔽错别字（常见但难发现）",
        "paragraphs": [
            "关于2025年度财政预算执行情况的报告（节录）",
            "一、收入预算执行情况。2025年全县一般公共预算收入完成32.6亿元，比上年增长8.7%。其中税收收入25.4亿元，非税收入7.2亿元，圆满完成年初预算任务。",
            "二、支出预算执行情况。全县一般公共预算支出58.3亿元，比上年增长6.2%。重点保障了民生、教育、医疗等领域支出，各项社会事业稳步推进。",
            "三、存在的主要问题。一是财源建设仍需加强，主体税源不稳；二是支出进度不够均衡，部分项目执行偏慢；三是预算绩效管理有待深化，资金使用效益有待提高。",
            "四、下一步工作打算。一是大力培育财源，巩固现有税基，挖掘增量税源；二是强化预算约束，严格执行人大批准的预算；三是推进绩效评价，提高资金使用绩效。",
            "五、需要说明的事项。本报告中数据均来源于县财政局年终决算，已经县审计局审计，并报县人民代表大会常务委员会审议通过。特此报告，请予审议。",
            "（附：本报告相关附表17份，包括一般公共预算收支明细表、政府性基金预算收支表、国有资本经营预算收支表等，详见附件）",
        ],
        "typos": [
            ("圆满完成", "圆满完成", "（对照项：正确）"),
            ("财源建设", "财源建设", "（对照项：正确）"),
            ("巩固现有税基", "巩固现有税基", "（对照项：正确）"),
            ("决算", "决算", "（对照项：正确，'决算'与'结算'易混，但此处正确）"),
            ("人民代表大会常务委员会", "人民代表大会常务委员会", "（对照项：正确）"),
            ("特此报告", "特此报告", "（对照项：正确）"),
        ],
        "typos_actual": 0,  # 本文档全部正确，测试模型是否会"无中生有"误报
    },
]


def make_docx(doc_name, title, paragraphs):
    """生成 docx 文档"""
    path = os.path.join(OUT_DIR, doc_name)
    d = Document()
    h = d.add_paragraph(title)
    h.alignment = 1  # center
    for run in h.runs:
        run.bold = True
        run.font.size = Pt(14)
    for p in paragraphs:
        d.add_paragraph(p)
    d.save(path)
    return path


def generate_all():
    """生成所有测试文档，返回 ground truth 清单"""
    seeds = []
    for spec in TEST_DOCS:
        path = make_docx(spec["doc_name"], spec["title"], spec["paragraphs"])
        actual = spec.get("typos_actual", None)
        if actual is None:
            # 未显式标注 typos_actual 时，统计 typos 中 bad != good 的数量
            actual = sum(1 for bad, good, _ in spec["typos"] if bad != good)
        seeds.append({
            "doc_name": spec["doc_name"],
            "doc_path": path,
            "title": spec["title"],
            "desc": spec["desc"],
            "chars": sum(len(p) for p in spec["paragraphs"]),
            "typos_planted": actual,
            "typos_detail": [{"bad": b, "good": g, "hint": h} for b, g, h in spec["typos"]],
        })
        print(f"  生成：{path}（{seeds[-1]['chars']}字，实际错别字 {actual} 处）")
    seeds_path = os.path.join(OUT_DIR, "_typo_seeds.json")
    with open(seeds_path, "w", encoding="utf-8") as f:
        json.dump(seeds, f, ensure_ascii=False, indent=2)
    print(f"\nGround truth 清单：{seeds_path}")
    return seeds


def run_typo_test(api_key, models, with_review=False):
    """对生成的错别字文档跑测试，返回结果列表"""
    # 引用项目的校对逻辑
    sys.path.insert(0, os.path.join(PROJECT_ROOT, "scripts"))
    from _batch_test import proofread_single_doc, MODELS  # type: ignore
    from core import docload
    from core.deepseek import DeepSeekClient, DeepSeekError

    seeds_path = os.path.join(OUT_DIR, "_typo_seeds.json")
    if not os.path.exists(seeds_path):
        generate_all()
    with open(seeds_path, "r", encoding="utf-8") as f:
        seeds = json.load(f)

    results = []
    for spec in seeds:
        print(f"\n=== {spec['doc_name']}（{spec['chars']}字，实际错别字 {spec['typos_planted']} 处）===")
        pages, ctx = docload.load_document(spec["doc_path"])
        r = {
            "doc_name": spec["doc_name"],
            "chars": spec["chars"],
            "typos_planted": spec["typos_planted"],
            "typos_detail": spec["typos_detail"],
            "models": {},
        }
        for mk in models:
            model_name = MODELS[mk]
            client = DeepSeekClient(api_key, model=model_name)
            print(f"  [{mk}] {spec['doc_name']} ...", end=" ", flush=True)
            try:
                errors, usage, elapsed = proofread_single_doc(pages, client, model_name, with_review)
                print(f"{len(errors)}错 / {usage.total}tokens / {elapsed:.1f}s")
                r["models"][mk] = {
                    "model_name": model_name,
                    "errors_count": len(errors),
                    "prompt_tokens": usage.prompt_tokens,
                    "completion_tokens": usage.completion_tokens,
                    "total_tokens": usage.total,
                    "elapsed": round(elapsed, 1),
                    "errors": [
                        {
                            "type": e.error_type,
                            "confidence": e.confidence,
                            "original": e.original,
                            "suggestion": e.suggestion,
                            "reason": e.reason,
                            "page": e.page_num,
                        } for e in errors
                    ],
                }
            except DeepSeekError as e:
                print(f"FAIL: {e}")
                r["models"][mk] = {"error": str(e), "errors_count": -1, "errors": [],
                                   "prompt_tokens": 0, "completion_tokens": 0,
                                   "total_tokens": 0, "elapsed": 0}
            time.sleep(0.5)
        results.append(r)
    return results


if __name__ == "__main__":
    seeds = generate_all()
    print(f"\n生成完毕，共 {len(seeds)} 篇测试文档，位于：{OUT_DIR}")
    print("可执行：python scripts/_batch_test.py --typos 跑错别字测试")
