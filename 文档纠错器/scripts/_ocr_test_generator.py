# -*- coding: utf-8 -*-
"""生成含典型 OCR 错误的测试文档，并跑 OCR 识别纠错模式测试。

输出：
    test_output/_ocr_test/ocr_01_form_similar.docx  形近字误识
    test_output/_ocr_test/ocr_02_number.docx        数字误识
    test_output/_ocr_test/ocr_03_punct.docx         标点误识
    test_output/_ocr_test/ocr_04_omission.docx      漏字漏行
    test_output/_ocr_test/ocr_05_repeat_garble.docx 重复识别+乱码
    test_output/_ocr_test/_ocr_seeds.json           ground truth 清单
    test_output/_ocr_test/_ocr_report.txt           测试报告

用法：
    .venv/Scripts/python.exe scripts/_ocr_test_generator.py
    .venv/Scripts/python.exe scripts/_ocr_test_generator.py --no-gen  # 只跑测试不重新生成
"""
import argparse
import json
import os
import sys
import time

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_ROOT)
sys.stdout.reconfigure(encoding="utf-8")
sys.stderr.reconfigure(encoding="utf-8")

from docx import Document

OUT_DIR = os.path.join(PROJECT_ROOT, "test_output", "_ocr_test")
os.makedirs(OUT_DIR, exist_ok=True)

SEEDS_PATH = os.path.join(OUT_DIR, "_ocr_seeds.json")


# ---------- 测试文档内容 ----------
# 每篇文档：正确文本 + OCR 错误文本 + 错误清单（错字, 对字, 类型, 理由）
# OCR 错误文本由正确文本逐处替换得到，确保 ground truth 可控。

DOCS = [
    {
        "file": "ocr_01_form_similar.docx",
        "title": "形近字误识测试",
        "correct": (
            "奉贤县志卷七·经济志\n"
            "清末，奉贤县境内商户百余家。民国初年，商业尚未兴盛，"
            "然已经有商会之设，联络商情，维系市面。\n"
            "其时，地方人士往来频繁，舟楫便利，货物通达四乡。"
            "戊戌年间，曾有记述，谓奉贤为海滨要邑，商贾云集。\n"
            "抗战胜利后，百废待兴，市场逐渐复苏，商铺渐增。"
        ),
        # (wrong, right, type, reason)：wrong=OCR识别错误文字, right=正确文字
        # 生成时把 correct 里的 right 替换为 wrong
        "errors": [
            ("尚末", "尚未", "错别字", "形近字误识：'未'被识为'末'（首横长短）"),
            ("己经", "已经", "错别字", "形近字误识：'已'被识为'己'（封口与否）"),
            ("入士", "人士", "错别字", "形近字误识：'人'被识为'入'"),
            ("戌戌", "戊戌", "错别字", "形近字误识：'戊'被识为'戌'"),
        ],
    },
    {
        "file": "ocr_02_number.docx",
        "title": "数字误识测试",
        "correct": (
            "卷八·食货志\n"
            "光绪十年，全县耕地共 3,285 顷，其中水田 2,170 顷，旱田 1,115 顷。\n"
            "是年，征收地丁银 8,964 两，漕粮 12,500 石。另征杂税若干，"
            "悉数解省。\n"
            "至 1949 年，工农业总产值 1,562 万元，比 1948 年增长 1.8 倍。"
        ),
        "errors": [
            # 数字误识造成逻辑矛盾：1948 被识成 1949，与前文 1949 矛盾
            ("1949 年增长", "1948 年增长", "逻辑", "数字误识：'8'被识为'9'，导致'比1949年增长'与上文'至1949年'矛盾"),
            # 数字 3 与 8 形近
            ("8,285", "3,285", "错别字", "数字误识：'3'被识为'8'"),
            # 数字 6 与 8 形近
            ("8,964", "8,964", "错别字", "提示：检查 8,964 是否被误识（实际未改，测模型能否发现疑点）"),
        ],
    },
    {
        "file": "ocr_03_punct.docx",
        "title": "标点误识测试",
        "correct": (
            "卷十一·大事记\n"
            "光绪元年（1875 年），知县张文炳到任。\n"
            "光绪三年，大水，淹没农田数千亩，朝廷下旨赈济。\n"
            "光绪五年，建书院，名曰“文溪书院”，延请名儒主讲。\n"
            "光绪八年，地震。房屋多有损毁，幸无伤亡。\n"
            "光绪十年，修《奉贤县志》，越三年而成。"
        ),
        "errors": [
            # 句号被识为逗号
            ("到任，", "到任。", "标点", "句号被识为逗号（句子已结束）"),
            ("赈济，", "赈济。", "标点", "句号被识为逗号（句子已结束）"),
            # 右引号识为左引号（方向错误）
            ("“文溪书院“", "“文溪书院”", "标点", "右引号被识为左引号（方向错误）"),
            # 逗号被识为句号
            ("地震。", "地震。", "标点", "提示：检查此处标点（实际为句号正确，测模型是否误报）"),
        ],
    },
    {
        "file": "ocr_04_omission.docx",
        "title": "漏字漏行测试",
        "correct": (
            "编纂体会\n"
            "修志之难，首在搜罗文献，次在考订事实，再次在秉笔直书。\n"
            "兹就编纂过程中所遇数端，略述如下，以备后来者参考。\n"
            "其一，资料之取舍。志书重在纪实，凡有关民生利弊者，均应详载。"
            "其无关宏旨者，则从简从略。\n"
            "其二，文字之锤炼。志书文字贵在简洁雅驯，力避浮华冗赘。"
            "凡引用旧志原文，均注明出处，以示征信。"
        ),
        "errors": [
            # 漏句：把"，再次在秉笔直书"整句漏掉
            ("次在考订事实。", "次在考订事实，再次在秉笔直书。", "逻辑", "漏句：'再次在秉笔直书'整句被漏识，上下文断裂"),
            # 漏字：把"则从简从略"的"略"漏掉
            ("则从简从。", "则从简从略。", "错别字", "漏字：'略'字被漏识"),
        ],
    },
    {
        "file": "ocr_05_repeat_garble.docx",
        "title": "重复识别与乱码测试",
        "correct": (
            "卷十·水利志\n"
            "奉贤地处海滨，河道港汊密布，水利为农政之大端。\n"
            "宋元以来，历代皆有修浚之举。明代嘉靖年间，知县张鹏率民浚河二十里，"
            "溉田千顷，民受其惠。\n"
            "清代雍正年间，开浚金汇港，引黄浦江水入县境。每年冬春之交，"
            "征集民夫，疏浚河道，岁以为常。"
        ),
        "errors": [
            # 重复识别
            ("海滨海滨", "海滨", "错别字", "重复识别：'海滨'被重复识别两次"),
            ("率民率民", "率民", "错别字", "重复识别：'率民'被重复"),
            # 乱码字符
            ("港Σ", "港汊", "错别字", "乱码：'汊'被识为无意义字符'Σ'"),
            ("冬春乊交", "冬春之交", "错别字", "乱码：'之'被识为怪字符'乊'"),
        ],
    },
    # ---- 版面错乱类（结构性错误，直接给出 OCR 错乱文本） ----
    {
        "file": "ocr_06_column_interleave.docx",
        "title": "两栏文字无序穿插",
        "raw_ocr": (
            "清晨携卷赴园，石亭置青瓷茶盏，风拂柳丝垂岸，沸水冲开碧螺春，"
            "碎花落满青石阶，清香漫过雕花木栏，驻足观池中游鱼，闲坐听林间雀鸣，"
            "水波轻晃碎云影，尘烦尽数消散。\n"
            "廊边新抽嫩笋，檐角悬竹制风铃，雨珠坠瓦当叮咚作响，茶烟袅袅绕廊柱，"
            "拾阶缓步登小丘，指尖轻触微凉竹篱，远望远山笼薄雾，抿一口清茶回甘绵长。\n"
            "篱边栽种海棠数株，杯中浮起细白茶毫，晚风吹散薄暮雾气，枝头蝶翼沾湿露，"
            "静坐半日不觉时序流转，池畔浮萍随涟漪缓缓舒展。"
        ),
        "correct": (
            "【左栏】清晨携卷赴园，风拂柳丝垂岸，碎花落满青石阶，驻足观池中游鱼，"
            "水波轻晃碎云影，尘烦尽数消散。廊边新抽嫩笋，雨珠坠瓦当叮咚作响，"
            "拾阶缓步登小丘，远望远山笼薄雾。篱边栽种海棠数株，晚风吹散薄暮雾气，"
            "枝头蝶翼沾湿露，静坐半日不觉时序流转，池畔浮萍随涟漪缓缓舒展。\n"
            "【右栏】石亭置青瓷茶盏，沸水冲开碧螺春，清香漫过雕花木栏，闲坐听林间雀鸣。"
            "檐角悬竹制风铃，茶烟袅袅绕廊柱，指尖轻触微凉竹篱，抿一口清茶回甘绵长。"
            "杯中浮起细白茶毫。"
        ),
        "errors": [
            ("两栏穿插", "应分栏", "逻辑", "版式错位：两栏文字被 OCR 按行交错读出，"
             "左栏（游园观景）与右栏（茶事器物）内容穿插混杂，语义断裂。"
             "如'清晨携卷赴园'后应接'风拂柳丝垂岸'，却插入右栏'石亭置青瓷茶盏'。"),
        ],
    },
    {
        "file": "ocr_07_sentence_half.docx",
        "title": "句子漏一半（残句拼接）",
        "raw_ocr": (
            "卷三·建置志\n"
            "奉贤县治，初在青村。明洪武十九年，设青村守御千户所。永乐年间，\n"
            "因海防需要，遂移治于奉贤城。城周九里，设四门：东曰海宴，西曰肃清，\n"
            "南曰明照，北曰。城内设县署、学署、典史署，分理政务。\n"
            "嘉靖年间，倭寇侵扰，城墙多有毁损。知县张鹏倡议修缮，\n"
            "历时三载，用工。修缮后，城垣坚固，民赖以安。\n"
            "万历五年，知县张栋增建敌楼三座，分别置于四门之上，"
            "以便瞭望。自此，海寇远遁，"
        ),
        "correct": (
            "卷三·建置志\n"
            "奉贤县治，初在青村。明洪武十九年，设青村守御千户所。永乐年间，\n"
            "因海防需要，遂移治于奉贤城。城周九里，设四门：东曰海宴，西曰肃清，\n"
            "南曰明照，北曰镇海。城内设县署、学署、典史署，分理政务。\n"
            "嘉靖年间，倭寇侵扰，城墙多有毁坏。知县张鹏倡议修缮，\n"
            "历时三载，用工万余。修缮后，城垣坚固，民赖以安。\n"
            "万历五年，知县张栋增建敌楼三座，分别置于四门之上，"
            "以便瞭望。自此，海寇远遁，地方晏然。"
        ),
        "errors": [
            ("北曰。", "北曰镇海。", "逻辑", "句子漏一半：'北曰镇海'的城门名整段漏识，只剩'北曰'加句号，语义残缺"),
            ("用工。", "用工万余。", "逻辑", "句子漏一半：'用工万余'的数量词漏识，只剩'用工'加句号"),
            ("海寇远遁，", "海寇远遁，地方晏然。", "逻辑", "句子漏一半：末句'地方晏然'漏识，以逗号结尾戛然而止"),
        ],
    },
    {
        "file": "ocr_08_para_merge.docx",
        "title": "段落拼接混乱（无分段）",
        "raw_ocr": (
            "卷九·学校志奉贤县学，在县治东南。明洪武三年建，中为明伦堂，"
            "东西列两斋。正统间，知县李敏重修，增建号舍二十间。"
            "清雍正五年，奉贤改隶松江府，学额因之调整。文庙在县学东，"
            "建有大成殿、东西庑、戟门、棂星门。乾隆八年，知县张鹏展修大成殿，"
            "增置祭器。咸丰十年，毁于战火。同治四年，知县陈延溥重建。"
            "书院旧有文溪书院，在县治西。康熙间，知县张鹏创建。"
            "道光六年，移建于南门外。光绪三十年，改为学堂。"
            "义学凡三所：一在城内，一在南桥，一在青村。皆设于乾隆年间，"
            "后废。社学凡五所，散布各乡，久废。"
        ),
        "correct": (
            "卷九·学校志\n"
            "【县学】奉贤县学，在县治东南。明洪武三年建，中为明伦堂，东西列两斋。"
            "正统间，知县李敏重修，增建号舍二十间。清雍正五年，奉贤改隶松江府，"
            "学额因之调整。\n"
            "【文庙】文庙在县学东，建有大成殿、东西庑、戟门、棂星门。"
            "乾隆八年，知县张鹏展修大成殿，增置祭器。咸丰十年，毁于战火。"
            "同治四年，知县陈延溥重建。\n"
            "【书院】书院旧有文溪书院，在县治西。康熙间，知县张鹏创建。"
            "道光六年，移建于南门外。光绪三十年，改为学堂。\n"
            "【义学】义学凡三所：一在城内，一在南桥，一在青村。皆设于乾隆年间，后废。\n"
            "【社学】社学凡五所，散布各乡，久废。"
        ),
        "errors": [
            ("段落拼接", "应分段", "逻辑", "版式错位：五个独立条目（县学/文庙/书院/义学/社学）的段落"
             "被 OCR 拼接成一大段，无分段区分，条目之间没有换行，导致内容归属混乱。"
             "如'增置祭器'后应另起一段开始'文庙'内容，却被直接拼到'咸丰十年'前。"),
        ],
    },
]


def generate_docs():
    """生成 8 个 OCR 错误测试 docx + ground truth json"""
    seeds = []
    for d in DOCS:
        planted = []
        if "raw_ocr" in d:
            # 版面错乱类：直接用 raw_ocr 作为文档内容，errors 为描述性
            ocr_text = d["raw_ocr"]
            for wrong, right, etype, reason in d["errors"]:
                planted.append({"ocr": wrong, "correct": right, "type": etype, "reason": reason})
        else:
            # 替换类：从 correct 按 errors 替换得到 OCR 错误文本
            ocr_text = d["correct"]
            for wrong, right, etype, reason in d["errors"]:
                if right in ocr_text and right != wrong:
                    ocr_text = ocr_text.replace(right, wrong, 1)
                planted.append({"ocr": wrong, "correct": right, "type": etype, "reason": reason})

        path = os.path.join(OUT_DIR, d["file"])
        doc = Document()
        doc.add_heading(d["title"], level=1)
        for para in ocr_text.split("\n"):
            if para.strip():
                doc.add_paragraph(para)
        doc.save(path)
        seeds.append({
            "file": d["file"],
            "title": d["title"],
            "chars": len(ocr_text),
            "planted": planted,
        })
        print(f"  生成 {d['file']}（{len(ocr_text)}字，{len(planted)}处 OCR 错误）")

    with open(SEEDS_PATH, "w", encoding="utf-8") as f:
        json.dump(seeds, f, ensure_ascii=False, indent=2)
    print(f"  ground truth 写入 {SEEDS_PATH}")


def run_test(api_key, model_name, with_review=False):
    """对 5 个测试文档跑 OCR 纠错模式"""
    from core import docload
    from core.deepseek import DeepSeekClient, DeepSeekError
    from core.proofread import proofread
    from core.models import TokenUsage

    client = DeepSeekClient(api_key, model=model_name)
    results = []
    for seed in seeds_global:
        path = os.path.join(OUT_DIR, seed["file"])
        try:
            pages, ctx = docload.load_document(path)
        except Exception as e:
            print(f"  !! 加载失败 {seed['file']}: {e}")
            continue

        print(f"  [{model_name}] {seed['file']} ({seed['chars']}字) ...", end=" ", flush=True)
        t0 = time.time()
        try:
            errors, usage = proofread(
                pages, client,
                token_limit=1000000,
                review=with_review,
                review_context=800,
                ocr_mode=True,
            )
            elapsed = time.time() - t0
            print(f"{len(errors)}错 / {usage.total}tokens / {elapsed:.1f}s")
            results.append({
                "file": seed["file"],
                "title": seed["title"],
                "chars": seed["chars"],
                "model": model_name,
                "planted": len(seed["planted"]),
                "errors_count": len(errors),
                "prompt_tokens": usage.prompt_tokens,
                "completion_tokens": usage.completion_tokens,
                "total_tokens": usage.total,
                "elapsed": round(elapsed, 1),
                "errors": [
                    {
                        "type": e.error_type,
                        "confidence": e.confidence,
                        "original": e.original,
                        "suggestion": e.suggestion,
                        "reason": e.reason,
                    } for e in errors
                ],
                "planted_list": seed["planted"],
            })
        except DeepSeekError as e:
            print(f"FAIL: {e}")
            results.append({
                "file": seed["file"], "model": model_name,
                "planted": len(seed["planted"]),
                "errors_count": -1, "total_tokens": 0, "elapsed": 0,
                "errors": [], "planted_list": seed["planted"], "error": str(e),
            })
        time.sleep(0.5)
    return results


def write_report(all_results):
    report_path = os.path.join(OUT_DIR, "_ocr_report.txt")
    lines = ["=" * 70,
             "DocProof OCR 识别纠错模式测试报告",
             f"测试时间：{time.strftime('%Y-%m-%d %H:%M:%S')}",
             "=" * 70]
    for r in all_results:
        lines.append(f"\n### {r['file']} - {r['title']}")
        lines.append(f"  模型：{r['model']}，字数：{r.get('chars', '?')}，"
                     f"埋设：{r['planted']}处，检出：{r['errors_count']}条")
        if r["errors_count"] >= 0:
            lines.append(f"  tokens：{r['total_tokens']}（p={r['prompt_tokens']}, "
                         f"c={r['completion_tokens']}），耗时：{r['elapsed']}s")
        lines.append(f"\n  【埋设的 OCR 错误（ground truth）】")
        for p in r["planted_list"]:
            lines.append(f"    - [{p['type']}] {p['ocr']!r} → 应为 {p['correct']!r}")
            lines.append(f"      {p['reason']}")
        lines.append(f"\n  【模型检出】")
        if r["errors_count"] <= 0:
            lines.append("    （无检出或失败）")
        else:
            for e in r["errors"][:30]:
                lines.append(f"    - [{e['type']}/{e['confidence']}] "
                             f"{e['original']!r} → {e['suggestion']!r}")
                lines.append(f"      理由：{e['reason']}")
    with open(report_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"\n报告已写入：{report_path}")
    return report_path


# 全局 seeds（run_test 用）
seeds_global = []


def main():
    global seeds_global
    parser = argparse.ArgumentParser()
    parser.add_argument("--no-gen", action="store_true", help="跳过文档生成，只跑测试")
    parser.add_argument("--models", default="flash,pro", help="模型，逗号分隔")
    parser.add_argument("--no-review", action="store_true", help="关闭二次复核")
    args = parser.parse_args()

    if not args.no_gen:
        print("生成 OCR 错误测试文档...")
        generate_docs()

    with open(SEEDS_PATH, "r", encoding="utf-8") as f:
        seeds_global = json.load(f)

    import config
    cfg = config.load()
    api_key = cfg["deepseek_api_key"]
    if not api_key:
        print("未配置 DeepSeek API Key")
        return

    models = {"flash": "deepseek-v4-flash", "pro": "deepseek-v4-pro"}
    selected = [m.strip() for m in args.models.split(",") if m.strip() in models]
    if not selected:
        selected = ["flash"]
    with_review = not args.no_review

    all_results = []
    for mk in selected:
        mn = models[mk]
        print(f"\n{'='*50}\n模型：{mn}（复核：{'开' if with_review else '关'}）\n{'='*50}")
        rs = run_test(api_key, mn, with_review)
        all_results.extend(rs)

    report_path = write_report(all_results)
    # 打印关键结论
    print("\n" + "=" * 50)
    print("测试结论速览")
    print("=" * 50)
    for r in all_results:
        status = f"检出 {r['errors_count']}/{r['planted']}"
        if r["errors_count"] == 0 and r["planted"] == 0:
            status += "（无错文档，零误报=正确）"
        elif r["errors_count"] == 0 and r["planted"] > 0:
            status += "（全漏报）"
        print(f"  {r['model']:20s} {r['file']:35s} {status}")


if __name__ == "__main__":
    main()
