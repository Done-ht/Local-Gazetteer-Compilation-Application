# -*- coding: utf-8 -*-
"""输出导出：可搜索 PDF / DOCX / TXT。

- PDF：有底图的页 = 插入压缩图 + 隐形文字层（render_mode=3，可选中/搜索/复制）；
  无底图的页（如 docx 文本页）= 纯文本页。
- DOCX：源文件名标题 + 逐页识别文本。
- TXT：逐页文本。
"""
from __future__ import annotations

import io
import os
from typing import List, Optional

from ..utils.text_dehardwrap import dehard_wrap, merge_lines_to_paragraphs
from .ocr_xfyun import OcrPage

# 底图像素 -> PDF 点的换算（等同按 150 DPI 渲染）
_SCALE = 72.0 / 150.0

_CJK_FONT_CACHE: Optional[tuple] = None


def _get_cjk_font() -> Optional[tuple]:
    """返回 (fontfile, fontname)；找不到返回 None（回退内置 china-s）。"""
    global _CJK_FONT_CACHE
    if _CJK_FONT_CACHE is not None:
        return _CJK_FONT_CACHE if _CJK_FONT_CACHE != ("", "") else None
    for fontfile, fontname in (
        ("C:/Windows/Fonts/simsun.ttc", "simsun"),
        ("C:/Windows/Fonts/msyh.ttc", "msyh"),
        ("C:/Windows/Fonts/simhei.ttf", "simhei"),
    ):
        if os.path.exists(fontfile):
            _CJK_FONT_CACHE = (fontfile, fontname)
            return _CJK_FONT_CACHE
    _CJK_FONT_CACHE = ("", "")
    return None


def build_pdf(pages: List[OcrPage], filename: str) -> bytes:
    """生成可搜索 PDF（有底图页叠加隐形文字层）。"""
    try:
        import fitz
    except ImportError:
        raise ImportError("未安装 PyMuPDF(fitz)，无法导出 PDF；请运行 setup.bat 或 pip install -r requirements.txt")

    font = _get_cjk_font()
    font_kwargs = {"fontname": "china-s"}
    if font:
        font_kwargs = {"fontfile": font[0], "fontname": font[1]}

    out = fitz.open()
    for pg in pages:
        if pg.image_bytes:
            w, h = pg.width, pg.height
            rect = fitz.Rect(0, 0, w * _SCALE, h * _SCALE)
            page = out.new_page(width=rect.width, height=rect.height)
            page.insert_image(rect, stream=pg.image_bytes)
            for ln in pg.lines:
                if not ln.text or ln.bbox == (0, 0, 0, 0):
                    continue
                x0, y0, x1, y1 = ln.bbox
                size = max((y1 - y0) * _SCALE * 0.9, 6)
                try:
                    page.insert_text(
                        (x0 * _SCALE, y1 * _SCALE),
                        ln.text,
                        fontsize=size,
                        render_mode=3,  # 不可见但可搜索
                        **font_kwargs,
                    )
                except Exception:
                    pass
        else:
            # 无底图页：纯文本页（docx 文本页等）
            page = out.new_page(width=595, height=842)  # A4
            y = 60
            for ln in pg.text.split("\n"):
                if not ln.strip():
                    continue
                if y > 800:
                    page = out.new_page(width=595, height=842)
                    y = 60
                try:
                    page.insert_text((50, y), ln.strip(), fontsize=12, **font_kwargs)
                except Exception:
                    pass
                y += 18
    buf = io.BytesIO()
    out.save(buf)
    out.close()
    return buf.getvalue()


def build_docx(pages: List[OcrPage], filename: str) -> bytes:
    """生成 DOCX：标题 + 逐页识别文本（段内硬换行已合并）。"""
    from docx import Document

    doc = Document()
    doc.add_heading(os.path.basename(filename), level=1)
    doc.add_heading("识别文字", level=2)
    for pg in pages:
        if pg.text.strip():
            doc.add_paragraph(f"—— 第 {pg.page_num} 页 ——")
            # 合并段内硬换行为段落（保留日期/序号/话题标签/短标题/空行等正常换行）
            for para in merge_lines_to_paragraphs(pg.text.split("\n")):
                if para.strip():
                    doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_txt(pages: List[OcrPage]) -> str:
    """生成纯文本（段内硬换行已合并，保留段落/条目等正常换行）。"""
    parts = []
    for pg in pages:
        if pg.text.strip():
            parts.append(f"===== 第 {pg.page_num} 页 =====")
            # 仅对页内容去硬换行，页标记独占一行由 \n\n 分隔
            parts.append(dehard_wrap(pg.text.strip()))
    return "\n\n".join(parts)
